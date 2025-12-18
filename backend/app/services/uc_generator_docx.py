# backend/app/services/uc_generator_docx.py
"""
Utilization Certificate (UC) DOCX Generator
Generates GFR 12-A format UC as Word document
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
import tempfile
from typing import Dict, Any


class UCWordGenerator:
    """Generate UC in Word format (GFR 12-A)"""
    
    @staticmethod
    def generate(uc_data: Dict[str, Any]) -> str:
        """
        Generate UC as Word document
        
        Args:
            uc_data: Complete UC data dictionary
            
        Returns:
            Path to generated Word file
        """
        doc = Document()
        
        # Set default font
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(10)
        
        # Page 1: Main UC Certificate
        UCWordGenerator._add_page1_certificate(doc, uc_data)
        
        # Page break
        doc.add_page_break()
        
        # Page 2: Certification Statements
        UCWordGenerator._add_page2_certifications(doc, uc_data)
        
        # Page break
        doc.add_page_break()
        
        # Page 3: Statement of Expenditure
        UCWordGenerator._add_page3_soe(doc, uc_data)
        
        # Save to temp file
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        temp_path = temp_file.name
        temp_file.close()
        
        doc.save(temp_path)
        return temp_path
    
    @staticmethod
    def _add_page1_certificate(doc: Document, data: Dict[str, Any]):
        """Add Page 1 - Main UC Certificate"""
        project = data['project']
        grants = data['grants_received']
        expenditure = data['expenditure']
        
        # Title
        title = doc.add_heading('GFR 12-A', level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        subtitle = doc.add_paragraph('{{See Rule 238 (1)}}S')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Main title
        main_title = doc.add_heading(
            'FORM OF UTILIZATION CERTIFICATE (UC)', level=2
        )
        main_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Period and grant type
        period_text = doc.add_paragraph(
            f'UTILIZATION CERTIFICATE FOR THE PERIOD '
            f'({data["period_from"].strftime("%d.%m.%Y")} to '
            f'{data["period_to"].strftime("%d.%m.%Y")})'
        )
        period_text.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        grant_type = doc.add_paragraph(
            'in respect of Recurring\n'
            'GRANTS-IN-AID/SALARIES/General Component/Recurring / Non-Recurring\n'
            'Creation of CAPITAL ASSESTS'
        )
        grant_type.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()  # Spacing
        
        # Project details
        details = [
            ('1.', 'Name of the Scheme', project.get('scheme', 'R & D in IT Group')),
            ('2.', 'Title of the Project', project.get('title', 'N/A')),
            ('3.', 'Institute Name', project.get('technical_group_name', 'N/A')),
            ('4.', 'Name of Principal Investigator (PI)', f"Dr. {project.get('pi_name', 'N/A')}"),
            ('5.', 'Sanction Order No & Date', 
             f"{project.get('sanctioned_number', 'N/A')} & Date: {project.get('start_date', 'N/A')}"),
            ('6.', 'Whether recurring or non-recurring grants', 'Recurring & Non-Recurring'),
        ]
        
        for num, label, value in details:
            p = doc.add_paragraph()
            p.add_run(f'{num} {label}').bold = True
            p.add_run(f': {value}')
        
        # Grant position section - Use correct numbering and title
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.add_run('8. ').bold = True
        p.add_run('Details of grants received, expenditure incurred and closing balances: (Actuals)').bold = True
        
        # Create grants position table
        UCWordGenerator._create_grants_position_table(doc, data)
        
        doc.add_paragraph()
        
        # Component-wise breakdown - USING EXPENDITURE
        p = doc.add_paragraph('Component wise utilization of grants:')
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        UCWordGenerator._create_component_breakdown_table(doc, expenditure)
        
        # Closing balance details
        doc.add_paragraph()
        doc.add_paragraph('Details of grants position at the end of the year')
        details_end = [
            f"(i)   Cash in Hand/Bank: Rs. {UCWordGenerator._format_currency(data['closing_balance'])}",
            f"(ii)  Refunds, if any: Rs. {UCWordGenerator._format_currency(data['closing_balance'])}- "
            f"Payment Advice Number (C082434921852)",
            f"(iii) Balance (Carry forward to next financial year): Rs. 0/-"
        ]
        
        for detail in details_end:
            doc.add_paragraph(detail)
        
        # Signatures section - DESIGNATION ONLY
        doc.add_paragraph()
        doc.add_paragraph()
        
        sig_table = doc.add_table(rows=1, cols=3)
        sig_table.autofit = False
        sig_table.allow_autofit = False
        
        cells = sig_table.rows[0].cells
        
        # Signature 1 - Only designation
        cells[0].text = '\n\nSignature with Seal:\nChief Administrative and Accounts Officer'
        
        # Signature 2 - Only designation
        cells[1].text = '\n\nSignature with Seal:\nHead of the Organization'
        
        # Signature 3 - Only designation
        cells[2].text = '\n\nSignature of PI:\nPrincipal Investigator'
    
    @staticmethod
    def _add_page2_certifications(doc: Document, data: Dict[str, Any]):
        """Add Page 2 - Certification Statements"""
        
        doc.add_heading('GFR 12-A', level=1)
        
        doc.add_paragraph(
            'Certified that I have satisfied myself that the conditions on which grants were '
            'sanctioned have been duly fulfilled/are being fulfilled and that I have exercised '
            'following checks to see that the money has been actually utilized for the purpose '
            'which it was sanctioned:'
        )
        
        certifications = [
            ('(i)', 'The main accounts and other subsidiary accounts and registers (including '
                    'assets registers) are maintained as prescribed in the relevant Act/Rules/standing '
                    'instructions (mention the act/Rules) and have been duly audited by designated '
                    'auditors. The figures depicted above tally with the audited figures mentioned in '
                    'financial statements/accounts.'),
            
            ('(ii)', 'There exist internal controls for safeguarding of public funds/assets, watching '
                     'outcomes and achievements of physical targets against the financial inputs, '
                     'ensuring quality in asset creation etc. & the periodic evaluation of internal '
                     'controls is exercised to ensure their effectiveness.'),
            
            ('(iii)', 'To the best of our knowledge and belief, no transactions have been entered '
                      'that are in violation of relevant Act/Rules/standing instructions and scheme guidelines.'),
            
            ('(iv)', 'The responsibilities among the key functionaries for execution of the scheme '
                     'have been assigned in clear terms and are not general in nature.'),
            
            ('(v)', 'The benefits were extended to the intended beneficiaries and only such '
                    'areas/districts were covered where the scheme was intended to operate.'),
            
            ('(vi)', 'The expenditure on various components of the scheme was in the proportions '
                     'authorized as per the scheme guidelines and terms and conditions of the grants-in-aid.'),
            
            ('(vii)', 'It has been ensured that the physical and financial performance under .... etc) '
                      '(name of the scheme) has been according to the requirements, as prescribed in '
                      'the guidelines issued by Govt. of India and the performance/targets achieved '
                      'statement for the year to which the utilization of the fund resulted in outcomes '
                      'given at Annexure-I duly enclosed.'),
            
            ('(viii)', 'The utilization of the fund resulted in outcomes given at Annexure-II duly '
                       'enclosed (to be formulated by the Ministry/Department concerned as per their '
                       'requirements/specifications)'),
            
            ('(ix)', 'Details of various schemes executed by the agency through grants-in-aid received '
                     'from the same Ministry or from other Ministries is enclosed at Annexure-II (to be '
                     'formulated by the Ministry/Department concerned as per their requirements/specifications)')
        ]
        
        for num, text in certifications:
            p = doc.add_paragraph()
            p.add_run(num).bold = True
            p.add_run(f' {text}')
        
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Date and place
        today = datetime.now()
        doc.add_paragraph(f'Date: {today.strftime("%d.%m.%Y")}')
        doc.add_paragraph('Place: Chennai')
        
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Signatures - Only designations
        sig_table = doc.add_table(rows=1, cols=3)
        cells = sig_table.rows[0].cells
        
        cells[0].text = '\n\nSignature with Seal:\nChief Administrative and Accounts Officer'
        cells[1].text = '\n\nSignature with Seal:\nHead of the Organization'
        cells[2].text = '\n\nSignature of PI:\nPrincipal Investigator'
    
    @staticmethod
    def _add_page3_soe(doc: Document, data: Dict[str, Any]):
        """Add Page 3 - Statement of Expenditure"""
        
        project = data['project']
        expenditure = data['expenditure']
        budget = data.get('budget', {})
        
        doc.add_heading('STATEMENT OF EXPENDITURE (SoE)', level=1)
        
        period_text = (
            f'for the period of {data["period_from"].strftime("%d.%m.%Y")} to '
            f'{data["period_to"].strftime("%d.%m.%Y")}'
        )
        p = doc.add_paragraph(period_text)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        subtitle = doc.add_paragraph(
            'in respect of Recurring GRANTS-IN-AID/SALARIES/General Component/Recurring /\n'
            'Non-Recurring Creation of CAPITAL ASSESTS'
        )
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        # Project details table
        details_table = doc.add_table(rows=7, cols=3)  # Changed from 6 to 7 rows to add Interest
        details_table.style = 'Table Grid'
        
        details_data = [
            ('1.', 'Sanction Letter No: No.4(4)/2021-ITEA', '6.', 'Grant Received in each year:'),
            ('2.', f'Total Project Cost: Rs. {UCWordGenerator._format_currency(sum(budget.values()))}', 
             'a.', '1st Payment: Rs. 88,01,000/-'),
            ('3.', 'Sanctioned/ Revised project cost (If applicable): Rs. /-', 
             'b.', '2nd Payment: Rs. 12,50,000/-'),
            ('4.', f'Date of Commencement of Project: {project.get("start_date", "N/A")}', 
             'c.', '3rd Payment: Rs. 19,18,071/-'),
            ('', '', 'd.', '4th Payment: Rs. 21,60,000/-'),
            ('5.', 'Statement of Expenditure', 'e.', f'Total: Rs. {UCWordGenerator._format_currency(data["grants_received"]["total"])}'),
            ('', '', 'f.', 'Interest: Rs. 0/-'),  # Added Interest row
        ]
        
        for i, (col1, col2, col3, col4) in enumerate(details_data):
            row = details_table.rows[i]
            row.cells[0].text = f'{col1} {col2}'
            if col3:
                row.cells[2].text = f'{col3} {col4}'
        
        doc.add_paragraph()
        
        # Expenditure table
        UCWordGenerator._create_soe_table(doc, expenditure, budget)
        
        doc.add_paragraph()
        
        # Footer info
        doc.add_paragraph(f'Funds released so far: {UCWordGenerator._format_currency(data["grants_received"]["total"])}/-')
        doc.add_paragraph(f'Date of start of Project: {project.get("start_date", "N/A")}')
        doc.add_paragraph(f'Expect Date of Completion: {project.get("end_date", "N/A")}')
        
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Signatures - Only designations
        sig_table = doc.add_table(rows=1, cols=2)
        cells = sig_table.rows[0].cells
        
        cells[0].text = '\n\nSignature of PI:\nPrincipal Investigator'
        cells[1].text = '\n\nSignature and Seal of\nHead of the Institute:\nExecutive Director'
        
        # Notes
        doc.add_paragraph()
        notes = [
            '1. Expenditure under the sanctioned heads, at any point of time, should not exceed '
            'funds allocated under that head, without prior approval of DST. Figures in Column (vii) '
            'should not exceed corresponding figures in Column (iii)',
            
            '2. Utilization Certificate (Annexure III) for each financial year ending 31st March has '
            'to be enclosed, along with request for carry-forward permission to next year.'
        ]
        
        for note in notes:
            doc.add_paragraph(note, style='List Number')
    
    @staticmethod
    def _create_grants_position_table(doc: Document, data: Dict[str, Any]):
        """Create the grants position table for Page 1 - matching reference format exactly"""
        
        grants = data['grants_received']
        expenditure = data['expenditure']
        opening = data['opening_balance']
        project = data['project']
        
        # Main table with complex structure
        table = doc.add_table(rows=4, cols=7)
        table.style = 'Table Grid'
        
        # Row 1: Main headers
        row1 = table.rows[0]
        headers1 = [
            'Unspent Balances of Grants received in previous years [figure as at SI. No. 7 (iii)]',
            'Interest Earned thereon',
            'Interest deposited back to the Government',
            'Grant received during the year',
            'Total available funds (1+2-3+4)',
            'Expenditure incurred',
            'Closing Balances (5-6)'
        ]
        
        for i, header in enumerate(headers1):
            row1.cells[i].text = header
            UCWordGenerator._set_cell_background(row1.cells[i], 'D3D3D3')
            # Center align and bold
            for paragraph in row1.cells[i].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.bold = True
        
        # Row 2: Column numbers
        row2 = table.rows[1]
        row2.cells[0].text = '1'
        row2.cells[1].text = '2'
        row2.cells[2].text = '3'
        row2.cells[3].text = '4'
        row2.cells[4].text = '5'
        row2.cells[5].text = '6'
        row2.cells[6].text = '7'
        
        for cell in row2.cells:
            UCWordGenerator._set_cell_background(cell, 'E8E8E8')
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Row 3: Sub-headers for column 4 (Grant received)
        row3 = table.rows[2]
        row3.cells[0].text = ''
        row3.cells[1].text = ''
        row3.cells[2].text = ''
        
        # Column 4 sub-headers
        row3.cells[3].text = 'Sanction no.\n(i)'
        row3.cells[4].text = 'Date (ii)'
        row3.cells[5].text = 'Amount\n(iii)'
        row3.cells[6].text = ''
        
        # Actually, we need to properly structure this - let me recreate
        # The reference shows column 4 split into 3 sub-columns
        # But we have 7 main columns, so we need a different approach
        
        # Let's merge cells properly
        # Merge row 2-3 for columns 1,2,3,5,6,7
        table.rows[1].cells[0].merge(table.rows[2].cells[0])
        table.rows[1].cells[1].merge(table.rows[2].cells[1])
        table.rows[1].cells[2].merge(table.rows[2].cells[2])
        table.rows[1].cells[4].merge(table.rows[2].cells[4])
        table.rows[1].cells[5].merge(table.rows[2].cells[5])
        table.rows[1].cells[6].merge(table.rows[2].cells[6])
        
        # Column 4 gets sub-headers in row 3
        row3.cells[3].text = 'Sanction no. (i)\nDate (ii)\nAmount (iii)'
        row3.cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Row 4: Data
        row4 = table.rows[3]
        row4.cells[0].text = UCWordGenerator._format_currency(opening)
        row4.cells[1].text = '0.00'
        row4.cells[2].text = '0.00'
        
        # Grant received cell with sanction details
        sanction_no = project.get('sanctioned_number', 'N/A')
        sanction_date = str(project.get('start_date', 'N/A'))
        grant_text = (
            f'{sanction_no}\n'
            f'{sanction_date}\n'
            f'{UCWordGenerator._format_currency(grants["total"])}'
        )
        row4.cells[3].text = grant_text
        row4.cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        total_available = opening + grants['total']
        row4.cells[4].text = UCWordGenerator._format_currency(total_available)
        row4.cells[5].text = UCWordGenerator._format_currency(expenditure['total'])
        row4.cells[6].text = UCWordGenerator._format_currency(data['closing_balance'])
        
        # Center align data cells
        for cell in row4.cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    @staticmethod
    def _create_component_breakdown_table(doc: Document, expenditure: Dict[str, Any]):
        """Create component-wise breakdown table using EXPENDITURE data"""
        
        # Calculate component-wise expenditure
        # General = Consumables + Travels + Contingencies + Overheads
        general_exp = (
            expenditure['by_head'].get('consumables', 0) +
            expenditure['by_head'].get('travels', 0) +
            expenditure['by_head'].get('contingencies', 0) +
            expenditure['by_head'].get('overheads', 0)
        )
        
        # Salary = Salaries (Recurring)
        salary_exp = expenditure['by_head'].get('salaries', 0)
        
        # Capital Assets = Equipment (Non-Recurring)
        capital_exp = expenditure['by_head'].get('equipment', 0)
        
        total_exp = general_exp + salary_exp + capital_exp
        
        table = doc.add_table(rows=2, cols=4)
        table.style = 'Table Grid'
        
        # Headers
        headers = ['Grant-in-aid-General', 'Grant-in-aid-Salary', 
                   'Grant-in-aid-creation of Capital Assets', 'Total']
        row1 = table.rows[0]
        for i, header in enumerate(headers):
            row1.cells[i].text = header
            UCWordGenerator._set_cell_background(row1.cells[i], 'D3D3D3')
        
        # Data - using expenditure amounts
        row2 = table.rows[1]
        row2.cells[0].text = UCWordGenerator._format_currency(general_exp)
        row2.cells[1].text = UCWordGenerator._format_currency(salary_exp)
        row2.cells[2].text = UCWordGenerator._format_currency(capital_exp)
        row2.cells[3].text = UCWordGenerator._format_currency(total_exp)
    
    @staticmethod
    def _create_soe_table(doc: Document, expenditure: Dict[str, Any], budget: Dict[str, float]):
        """Create Statement of Expenditure table"""
        
        table = doc.add_table(rows=8, cols=5)
        table.style = 'Table Grid'
        
        # Headers
        headers_row = table.rows[0]
        headers = ['SI. NO.', 'HEAD OF EXPENDITURE AS PER SANCTION ORDER', 
                   'Total Approved Budget', 'Expenditure Incurred 01.04.2024 to 09.03.2025',
                   'TOTAL EXPENDITURE']
        
        for i, header in enumerate(headers):
            headers_row.cells[i].text = header
            UCWordGenerator._set_cell_background(headers_row.cells[i], 'D3D3D3')
        
        # Data rows
        heads_data = [
            ('1', 'Equipment', 'equipment'),
            ('2', 'Salaries', 'salaries'),
            ('3', 'Consumables', 'consumables'),
            ('4', 'Travels', 'travels'),
            ('5', 'Contingencies', 'contingencies'),
            ('6', 'Overheads', 'overheads')
        ]
        
        total_budget = 0
        total_exp = 0
        
        for i, (num, label, key) in enumerate(heads_data, start=1):
            row = table.rows[i]
            row.cells[0].text = num
            row.cells[1].text = label
            
            budget_amt = budget.get(key, 0)
            exp_amt = expenditure['by_head'].get(key, 0)
            
            row.cells[2].text = UCWordGenerator._format_currency(budget_amt)
            row.cells[3].text = UCWordGenerator._format_currency(exp_amt)
            row.cells[4].text = UCWordGenerator._format_currency(exp_amt)
            
            total_budget += budget_amt
            total_exp += exp_amt
        
        # Total row
        total_row = table.rows[7]
        total_row.cells[0].text = ''
        total_row.cells[1].text = 'Total'
        total_row.cells[2].text = UCWordGenerator._format_currency(total_budget)
        total_row.cells[3].text = UCWordGenerator._format_currency(total_exp)
        total_row.cells[4].text = UCWordGenerator._format_currency(total_exp)
        
        for cell in total_row.cells:
            UCWordGenerator._set_cell_background(cell, 'E8E8E8')
    
    @staticmethod
    def _format_currency(amount: float) -> str:
        """Format currency in Indian style"""
        if amount == 0:
            return "0.00"
        
        amount_str = f"{amount:,.2f}"
        return f"Rs. {amount_str}"
    
    @staticmethod
    def _set_cell_background(cell, color_hex: str):
        """Set cell background color"""
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), color_hex)
        cell._element.get_or_add_tcPr().append(shading_elm)